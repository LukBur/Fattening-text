from docx import Document
from docx.shared import Pt, RGBColor
import os

doc = Document("reduta.docx")

# bold every next max 3 letters of the word
def first_3(word, num, paragraph):
    for i, litera in enumerate(word):
        run = paragraph.add_run(litera)
        if i < num:
            run.bold = True
    paragraph.add_run(" ")

new_doc = Document()

# checking lines -> words split in lines
for para in doc.paragraphs:
    p = new_doc.add_paragraph() 
    words = para.text.split()
    for word in words:
        first_3(word, 3, p)

new_doc.save("reduta_enhanced.docx")

# showing the result
os.startfile("reduta_enhanced.docx")
msg = input()

#closing
os.system("taskkill /f /im WINWORD.EXE")
print(msg) # kill message, because why not